# infrastructure/document_loader.py
# 文档解析模块——Phase 2 的核心

from docx import Document                          # python-docx 的核心类
from docx.text.paragraph import Paragraph           # 段落类型，用于类型判断
from docx.table import Table                         # 表格类型，用于类型判断
from models.schemas import Section, TableData, ParsedDocument  # Step 2.4: 加入 ParsedDocument


# ===== Step 2.3 新增：提取表格数据 =====

def _table_to_text(table: Table) -> str:
    """
    把一个表格转成"纯文本版"，方便 LLM 阅读

    生活比喻：把 Excel 表格打印成一页纸——每行占一行，列之间用 | 分隔

    为什么需要这个？LLM 不认识表格对象，它只认识文本。
    我们把表格转成类似 Markdown 表格的格式，LLM 一看就懂。

    参数：
        table: python-docx 的表格对象

    返回：
        表格的文本表示，如：
        序号 | 设备名称 | 数量 | 单价
        1 | 服务器 | 10 | 50000
        2 | 交换机 | 5 | 8000
    """
    lines = []  # 存储每一行的文本

    for i, row in enumerate(table.rows):
        # 提取这一行每个单元格的文本，去掉首尾空格
        cells = [cell.text.strip() for cell in row.cells]
        # 用 " | " 把单元格内容拼起来
        line = " | ".join(cells)
        lines.append(line)

        # 第一行（表头）后面加一条分隔线，类似 Markdown 表格
        if i == 0:
            # 生成 "--- | --- | ---" 这样的分隔线
            separator = " | ".join(["---"] * len(cells))
            lines.append(separator)

    return "\n".join(lines)


def extract_tables(file_path: str) -> list[TableData]:
    """
    从 Word 文档中提取所有表格

    生活比喻：把菜单上所有的价目表都抄下来，注明来源在"第几页"

    为什么不用 Document.tables？
    因为 Document.tables 只给你表格对象，不告诉你它在文档的哪个位置。
    我们需要遍历文档的 body element，才能知道表格出现在哪个章节附近。

    这里用一个简化方案：遍历 Document.tables，同时记录表格序号
    后面在 extract_sections 中，我们会把表格"挂"到对应的章节上。

    参数：
        file_path: Word 文件的路径

    返回：
        TableData 列表
    """
    doc = Document(file_path)
    tables_data: list[TableData] = []

    for i, table in enumerate(doc.tables):
        # 1. 提取表头（第一行）
        headers = []
        if table.rows:
            headers = [cell.text.strip() for cell in table.rows[0].cells]

        # 2. 提取数据行（第二行开始）
        rows = []
        for row in table.rows[1:]:
            row_data = [cell.text.strip() for cell in row.cells]
            rows.append(row_data)

        # 3. 生成文本版
        as_text = _table_to_text(table)

        # 4. 暂时用表格序号作为 location，后面会修正
        location = f"表格 #{i + 1}"

        # 5. 组装成 TableData 对象
        table_data = TableData(
            headers=headers,
            rows=rows,
            as_text=as_text,
            location=location
        )
        tables_data.append(table_data)

    return tables_data


def extract_tables_with_context(file_path: str) -> list[TableData]:
    """
    提取表格，并尝试推断每个表格所在的章节位置

    生活比喻：不仅抄了价目表，还标注了"这张表在菜单的哪个板块下面"

    核心思路：遍历文档的所有 body 元素（段落和表格交替出现），
    维护一个"当前最近的标题"变量，遇到表格时就把当前标题作为它的 location。

    参数：
        file_path: Word 文件的路径

    返回：
        带有 location 信息的 TableData 列表
    """
    doc = Document(file_path)
    tables_data: list[TableData] = []

    # 当前最近的标题，用来标记表格属于哪个章节
    current_heading = "文档开头"

    # 表格索引——因为 doc.tables 是纯表格列表，
    # 我们需要知道当前处理的是第几个表格
    table_index = 0

    # 遍历文档 body 的所有子元素
    # doc.element.body 是底层的 XML 元素
    # iterchildren() 按顺序遍历所有子节点（段落和表格交替出现）
    for child in doc.element.body.iterchildren():
        # 判断这个子元素是段落还是表格
        # python-docx 用不同的 XML tag 区分：
        #   <w:p> = 段落（paragraph）
        #   <w:tbl> = 表格（table）
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag

        if tag == 'p':
            # 这是一个段落——检查它是不是标题
            # 需要从 XML 元素还原回 Paragraph 对象
            # python-docx 提供了 CT_P → Paragraph 的方式
            para = Paragraph(child, doc)
            text = para.text.strip()

            if text and _get_heading_level(para) is not None:
                # 是标题，更新 current_heading
                current_heading = text

        elif tag == 'tbl':
            # 这是一个表格——提取数据
            if table_index < len(doc.tables):
                table = doc.tables[table_index]

                # 提取表头
                headers = []
                if table.rows:
                    headers = [cell.text.strip() for cell in table.rows[0].cells]

                # 提取数据行
                rows = []
                for row in table.rows[1:]:
                    row_data = [cell.text.strip() for cell in row.cells]
                    rows.append(row_data)

                # 生成文本版
                as_text = _table_to_text(table)

                # 用当前标题作为 location
                location = current_heading

                table_data = TableData(
                    headers=headers,
                    rows=rows,
                    as_text=as_text,
                    location=location
                )
                tables_data.append(table_data)

                table_index += 1

    return tables_data


# ===== Step 2.1 的函数：读取所有段落 =====

def read_paragraphs(file_path: str) -> list[str]:
    """
    读取 Word 文件的所有段落文本（Step 2.1 已完成）

    参数：
        file_path: Word 文件的路径，如 "data/input/标书.docx"

    返回：
        所有段落文本的列表
    """
    doc = Document(file_path)
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)
    return paragraphs


# ===== Step 2.2 新增：提取章节结构 =====

def _get_heading_level(paragraph: Paragraph) -> int | None:
    """
    判断一个段落是否是标题，如果是返回标题层级（1/2/3），不是返回 None

    生活比喻：检查菜单上的每一行，判断它是"大分类"、"中分类"还是"小分类"

    参数：
        paragraph: python-docx 的段落对象

    返回：
        标题层级（1/2/3）或 None（表示是正文）
    """
    # Word 的标题样式名通常是 "Heading 1"、"Heading 2"、"Heading 3"
    style_name = paragraph.style.name

    # 检查样式名是否以 "Heading" 开头
    if style_name.startswith("Heading"):
        # 提取层级数字，如 "Heading 1" → 1
        try:
            level = int(style_name.split()[-1])
            return level
        except ValueError:
            # 如果样式名是 "Heading" 但没有数字（罕见情况），返回 None
            return None

    return None  # 不是标题


def extract_sections(file_path: str) -> list[Section]:
    """
    从 Word 文档中提取章节结构（树状）

    生活比喻：把一本平铺的菜单，整理成"大分类 → 中分类 → 小分类"的目录树

    核心思路：
    1. 遍历文档的每一个段落
    2. 判断它是标题还是正文
    3. 如果是标题 → 创建一个新的 Section 节点
    4. 如果是正文 → 追加到当前最近的 Section 的 content 中
    5. 最后根据标题层级组装成树状结构

    参数：
        file_path: Word 文件的路径

    返回：
        顶层 Section 列表（每个 Section 可能有 children 子章节）
    """
    doc = Document(file_path)

    # ---- 第一步：线性扫描，把每个段落归类到对应的 Section ----
    # flat_sections 存储所有 Section（平铺列表，还没组装树）
    flat_sections: list[Section] = []
    current_section: Section | None = None  # 当前正在"接收正文"的章节

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue  # 跳过空段落

        heading_level = _get_heading_level(para)

        if heading_level is not None:
            # ---- 这是一个标题 ----
            # 创建新的 Section 节点
            new_section = Section(
                title=text,                       # 标题文本
                level=heading_level,               # 标题层级 1/2/3
                content="",                        # 正文先留空，后面往里填
                tables=[],                         # 表格后面再处理（Step 2.3）
                page_range=(0, 0),                 # 页码后面再处理
                children=[]                        # 子章节，后面组装树时填充
            )
            flat_sections.append(new_section)
            current_section = new_section          # 切换"当前章节"指针

        else:
            # ---- 这是正文 ----
            if current_section is not None:
                # 把这段正文追加到当前章节的 content 中
                if current_section.content:
                    current_section.content += "\n" + text
                else:
                    current_section.content = text
            # 如果 current_section 是 None，说明正文出现在第一个标题之前
            # 这种情况我们暂时忽略（通常是封面或目录）

    # ---- 第二步：组装树状结构 ----
    # 核心思路：用一个"栈"来维护当前的层级关系
    # 栈顶永远是"当前正在处理的父节点"
    #
    # 生活比喻：你在整理书架
    # - 看到"第一章"（Level 1）→ 放到书架第一层
    # - 看到"1.1 节"（Level 2）→ 放到"第一章"下面
    # - 看到"1.1.1 小节"（Level 3）→ 放到"1.1 节"下面
    # - 看到"第二章"（Level 1）→ 回到书架第一层，新开一个位置

    root_sections: list[Section] = []   # 最终的顶层结果
    stack: list[Section] = []           # 维护层级关系的栈

    for section in flat_sections:
        # 弹出栈中所有层级 >= 当前层级的节点
        # （因为它们不是当前节点的父节点）
        while stack and stack[-1].level >= section.level:
            stack.pop()

        if stack:
            # 栈不为空 → 当前节点是栈顶节点的子节点
            stack[-1].children.append(section)
        else:
            # 栈为空 → 当前节点是顶层节点
            root_sections.append(section)

        # 把当前节点压入栈
        stack.append(section)

    return root_sections


def print_section_tree(sections: list[Section], indent: int = 0) -> None:
    """
    以树状格式打印章节结构（调试用）

    生活比喻：打印一份漂亮的目录，让你一眼看出层级关系

    参数：
        sections: Section 列表
        indent: 当前缩进层级（用于递归）
    """
    for section in sections:
        prefix = "  " * indent  # 缩进，每层加两个空格
        # 只显示前 60 个字符，避免标题太长
        title_short = section.title[:60] + "..." if len(section.title) > 60 else section.title
        print(f"{prefix}📌 [{section.level}] {title_short}")
        # 递归打印子章节
        print_section_tree(section.children, indent + 1)


# ===== Step 2.4 新增：DocumentLoader 类（一键解析）=====

def _collect_full_text(sections: list[Section]) -> str:
    """
    递归收集所有章节的正文文本，拼成一份完整的全文

    生活比喻：把目录树里每一页的内容按顺序抄到一张长纸上

    为什么需要？后面 LLM 做行业分析时需要看全文，
    而且 ParsedDocument 也需要 full_text 字段。

    参数：
        sections: 章节树（可能有嵌套的 children）

    返回：
        拼接后的全文文本
    """
    parts = []
    for section in sections:
        # 加上这个章节的正文
        if section.content.strip():
            parts.append(section.content)
        # 递归收集子章节的正文
        if section.children:
            parts.append(_collect_full_text(section.children))
    return "\n".join(parts)


def _count_sections(sections: list[Section]) -> int:
    """
    递归统计章节总数（包括子章节）

    生活比喻：数一数目录树里一共有多少个节点
    """
    count = 0
    for section in sections:
        count += 1  # 当前章节
        if section.children:
            count += _count_sections(section.children)  # 子章节
    return count


def _attach_tables_to_sections(sections: list[Section], tables: list[TableData]) -> None:
    """
    把表格"挂"到对应的章节上

    生活比喻：把抄好的价目表放回菜单对应的板块下面

    核心逻辑：
    - 遍历每个表格，看它的 location 是否包含某个章节的 title
    - 如果匹配上了，就把这个表格加到那个章节的 tables 列表里
    - 匹配策略：location 包含 section.title，或 section.title 包含 location

    参数：
        sections: 章节树
        tables: 带 location 的表格列表
    """
    for table in tables:
        # 递归查找匹配的章节
        _find_and_attach(sections, table)


def _find_and_attach(sections: list[Section], table: TableData) -> bool:
    """
    递归查找并把表格挂到匹配的章节上

    返回：True 表示找到了匹配的章节，False 表示没找到
    """
    for section in sections:
        # 匹配策略：双向包含检查
        # 因为 location 可能是"第六章 报价清单"，而 section.title 是"报价清单"
        if (table.location in section.title or
                section.title in table.location or
                table.location == section.title):
            section.tables.append(table)
            return True

        # 递归查找子章节
        if section.children:
            if _find_and_attach(section.children, table):
                return True

    return False


class DocumentLoader:
    """
    文档加载器——一键把 Word 文件变成 ParsedDocument

    生活比喻：全自动洗碗机——你把脏碗放进去，按一个按钮，
    出来就是干净的碗。你不需要管里面经历了冲洗、烘干、消毒。

    用法：
        loader = DocumentLoader()
        doc = loader.load("data/input/标书.docx")
        print(doc.filename)          # 文件名
        print(len(doc.sections))     # 章节数
        print(doc.full_text[:200])   # 全文前 200 字
    """

    def load(self, file_path: str) -> ParsedDocument:
        """
        加载 Word 文件，返回完整的 ParsedDocument

        内部流程：
        1. 提取章节结构（extract_sections）
        2. 提取表格并定位（extract_tables_with_context）
        3. 把表格挂到对应章节上
        4. 收集全文文本
        5. 组装成 ParsedDocument

        参数：
            file_path: Word 文件路径

        返回：
            ParsedDocument 对象
        """
        import os

        # 第一步：提取章节结构
        sections = extract_sections(file_path)

        # 第二步：提取表格（带位置信息）
        tables = extract_tables_with_context(file_path)

        # 第三步：把表格挂到对应章节上
        _attach_tables_to_sections(sections, tables)

        # 第四步：收集全文文本
        full_text = _collect_full_text(sections)

        # 第五步：统计信息
        total_sections = _count_sections(sections)
        filename = os.path.basename(file_path)

        # 第六步：组装成 ParsedDocument
        parsed_doc = ParsedDocument(
            filename=filename,
            total_pages=0,  # python-docx 无法直接获取页数，暂设为 0
            sections=sections,
            full_text=full_text,
            metadata={
                "file_path": file_path,
                "total_sections": total_sections,
                "total_tables": len(tables),
                "total_chars": len(full_text),
            }
        )

        return parsed_doc


# ===== 测试代码 =====
if __name__ == "__main__":
    file_path = "data/input/【定稿0424】湖北省医疗器械综合监管建设.docx"

    print("=" * 60)
    print("Step 2.2 测试：提取章节结构")
    print("=" * 60)

    sections = extract_sections(file_path)

    print(f"\n共找到 {len(sections)} 个顶层章节：\n")
    print_section_tree(sections)

    print("\n" + "=" * 60)
    print("Step 2.3 测试：提取表格数据")
    print("=" * 60)

    # 测试基础表格提取
    tables = extract_tables(file_path)
    print(f"\n共找到 {len(tables)} 个表格：\n")
    for i, table in enumerate(tables):
        print(f"--- 表格 #{i + 1} ---")
        print(f"表头: {table.headers}")
        print(f"数据行数: {len(table.rows)}")
        print(f"文本预览:\n{table.as_text[:300]}...")
        print()

    # 测试带章节定位的表格提取
    print("=" * 60)
    print("Step 2.3 进阶测试：带章节定位的表格提取")
    print("=" * 60)

    tables_with_ctx = extract_tables_with_context(file_path)
    print(f"\n共找到 {len(tables_with_ctx)} 个表格（带位置信息）：\n")
    for i, table in enumerate(tables_with_ctx):
        print(f"  表格 #{i + 1} | 所在章节: {table.location} | 表头: {table.headers[:3]}...")

    print("\n" + "=" * 60)
    print("Step 2.4 测试：DocumentLoader 一键解析")
    print("=" * 60)

    loader = DocumentLoader()
    doc = loader.load(file_path)

    print(f"\n📄 文件名: {doc.filename}")
    print(f"📊 总章节数: {doc.metadata['total_sections']}")
    print(f"📋 总表格数: {doc.metadata['total_tables']}")
    print(f"📝 全文字符数: {doc.metadata['total_chars']}")
    print(f"\n🌳 章节树（前 5 个顶层章节）：")
    for section in doc.sections[:5]:
        tables_count = len(section.tables)
        print(f"  [{section.level}] {section.title[:50]}  （含 {tables_count} 个表格）")
        for child in section.children[:3]:
            child_tables = len(child.tables)
            print(f"      [{child.level}] {child.title[:50]}  （含 {child_tables} 个表格）")

    print(f"\n📖 全文预览（前 200 字）：")
    print(doc.full_text[:200])
