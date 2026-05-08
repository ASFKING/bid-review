# generate_test_docx.py
# 生成一份测试用的投标文件（.docx）
# 运行：python generate_test_docx.py
# 输出：data/input/测试投标文件.docx

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path


def create_test_bid_document():
    """创建一份故意埋了问题的投标文件"""

    doc = Document()

    # ===== 封面 =====
    doc.add_paragraph("")
    doc.add_paragraph("")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("XX市智慧城市大数据平台建设项目")
    run.font.size = Pt(22)
    run.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("投 标 文 件")
    run.font.size = Pt(28)
    run.bold = True

    doc.add_paragraph("")
    doc.add_paragraph("")

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run("投标人：深圳某科技有限公司")
    run.font.size = Pt(14)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run("2026年4月")
    run.font.size = Pt(14)

    doc.add_page_break()

    # ===== 第一章 投标函 =====
    doc.add_heading("第一章 投标函", level=1)
    doc.add_paragraph(
        "致：XX市大数据管理局\n\n"
        "我方深圳某科技有限公司（以下简称“投标人”）已仔细阅读并充分理解"
        "《XX市智慧城市大数据平台建设项目招标文件》（招标编号：ZB-2026-0415）"
        "的全部内容，包括但不限于技术要求、商务条款、评标标准等。\n\n"
        "我方愿意按照招标文件规定的各项要求，以人民币壹佰捌拾万元整"
        "（¥1,800,000.00）的投标总价，承担本项目的全部工作。\n\n"
        "我方承诺：\n"
        "1. 本投标文件中所有信息真实、准确、完整；\n"
        "2. 我方完全接受招标文件的所有条款；\n"
        "3. 本投标文件有效期为自开标之日起90个日历日。"
    )
    doc.add_paragraph("")

    # ===== 第二章 公司概况 =====
    doc.add_heading("第二章 公司概况", level=1)

    doc.add_heading("2.1 公司简介", level=2)
    doc.add_paragraph(
        "深圳某科技有限公司成立于2015年，注册资本5000万元人民币，"
        "是一家专注于智慧城市、大数据平台建设的高新技术企业。"
        "公司总部位于深圳市南山区科技园，在北京、上海、广州设有分公司。\n\n"
        "公司现有员工200余人，其中技术人员占比75%，"
        "拥有博士学位15人、硕士学位60人。"
    )

    doc.add_heading("2.2 资质情况", level=2)
    doc.add_paragraph(
        "公司拥有以下主要资质：\n"
        "- 软件企业认定证书\n"
        "- ISO 9001质量管理体系认证\n"
        "- ISO 27001信息安全管理体系认证\n"
        "- CMMI L3认证"
    )

    doc.add_heading("2.3 类似项目经验", level=2)
    doc.add_paragraph(
        "近三年公司承担的主要类似项目如下：\n\n"
        "1. 某省会城市智慧交通平台建设项目（合同金额：2200万元）\n"
        "2. 某地级市政务数据中台建设项目（合同金额：1500万元）\n"
        "3. 某开发区物联网监控平台项目（合同金额：800万元）"
    )

    # ===== 第三章 技术方案 =====
    doc.add_heading("第三章 技术方案", level=1)

    doc.add_heading("3.1 总体架构设计", level=2)
    doc.add_paragraph(
        "本项目采用平台+应用的总体架构思路，分为基础设施层、数据资源层、"
        "应用支撑层和业务应用层四个层次。\n\n"
        "基础设施层采用混合云架构，核心数据部署在政务云，"
        "非敏感业务部署在公有云，通过专线互联。\n\n"
        "数据资源层建设统一的数据湖，汇聚各委办局数据资源，"
        "实现数据的采集、清洗、存储、治理全流程管理。"
    )

    doc.add_heading("3.2 数据采集方案", level=2)
    doc.add_paragraph(
        "数据采集采用实时+批量双模式：\n\n"
        "- 实时采集：基于 Kafka 消息队列，支持日均1000万条消息\n"
        "- 批量采集：基于 Apache Spark，支持每日TB级数据同步\n"
        "- API网关：统一的数据服务接口，支持RESTful和gRPC协议"
    )

    doc.add_heading("3.3 安全设计方案", level=2)
    doc.add_paragraph(
        "本项目安全设计遵循《网络安全等级保护基本要求》（GB/T 22239-2019）"
        "三级要求，主要措施包括：\n\n"
        "1. 网络安全：部署防火墙、入侵检测系统、WAF\n"
        "2. 数据安全：敏感数据加密存储，传输全程TLS加密\n"
        "3. 访问控制：基于RBAC的权限管理，支持多因素认证\n"
        "4. 审计日志：全操作审计，日志保留180天"
    )

    # 故意埋坑：复制粘贴痕迹
    doc.add_heading("3.4 系统集成方案", level=2)
    doc.add_paragraph(
        "本系统需要与以下外部系统对接：\n\n"
        "1. 某市政务服务平台（API对接）\n"
        "2. 某省公安信息查询系统（数据库对接）\n"
        "3. 某省水利监测系统（消息队列对接）\n\n"
        "对接方式采用ESB企业服务总线，实现各系统间的松耦合集成。"
    )

    # ===== 第四章 项目实施计划 =====
    doc.add_heading("第四章 项目实施计划", level=1)

    doc.add_heading("4.1 项目组织", level=2)
    doc.add_paragraph(
        "本项目设立项目管理办公室（PMO），由项目经理统一协调。"
        "项目团队配置如下：\n\n"
        "- 项目经理：1人（PMP认证，10年项目管理经验）\n"
        "- 技术总监：1人\n"
        "- 开发工程师：8人\n"
        "- 测试工程师：3人\n"
        "- 实施工程师：2人"
    )

    doc.add_heading("4.2 实施进度", level=2)
    table = doc.add_table(rows=6, cols=4)
    table.style = "Table Grid"
    headers = ["阶段", "工作内容", "工期（天）", "交付物"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    data = [
        ["需求分析", "需求调研、需求规格说明书", "30", "需求规格说明书"],
        ["系统设计", "概要设计、详细设计", "25", "设计文档"],
        ["编码开发", "功能开发、单元测试", "60", "源代码"],
        ["系统测试", "集成测试、性能测试、安全测试", "20", "测试报告"],
        ["部署上线", "系统部署、数据迁移、用户培训", "15", "验收报告"],
    ]
    for row_idx, row_data in enumerate(data):
        for col_idx, cell_data in enumerate(row_data):
            table.rows[row_idx + 1].cells[col_idx].text = cell_data
    doc.add_paragraph("")

    # 故意埋坑：内容太短
    doc.add_heading("4.3 风险管理", level=2)
    doc.add_paragraph("项目风险主要包括技术风险和进度风险，将通过定期会议跟踪。")

    # ===== 第五章 报价清单 =====
    doc.add_heading("第五章 报价清单", level=1)

    doc.add_heading("5.1 报价汇总表", level=2)
    summary_table = doc.add_table(rows=5, cols=3)
    summary_table.style = "Table Grid"
    headers = ["序号", "费用项目", "金额（万元）"]
    for i, h in enumerate(headers):
        summary_table.rows[0].cells[i].text = h
    summary_data = [
        ["1", "软件开发费", "120"],
        ["2", "硬件采购费", "35"],
        ["3", "实施服务费", "15"],
        ["4", "合计", "170"],  # ← 和投标函的180万不一致！
    ]
    for row_idx, row_data in enumerate(summary_data):
        for col_idx, cell_data in enumerate(row_data):
            summary_table.rows[row_idx + 1].cells[col_idx].text = cell_data
    doc.add_paragraph("")

    # 故意埋坑：只写"详见附件"
    doc.add_heading("5.2 软件开发费明细", level=2)
    doc.add_paragraph("详见附件。")

    # ===== 第六章 售后服务方案 =====
    doc.add_heading("第六章 售后服务方案", level=1)

    doc.add_heading("6.1 服务承诺", level=2)
    doc.add_paragraph(
        "我方承诺提供以下售后服务：\n\n"
        "1. 质保期：系统验收后12个月免费质保\n"
        "2. 响应时间：\n"
        "   - 紧急问题（系统瘫痪）：30分钟响应，4小时解决\n"
        "   - 严重问题（功能异常）：2小时响应，8小时解决\n"
        "   - 一般问题：4小时响应，24小时解决\n"
        "3. 远程支持：7×24小时远程技术支持\n"
        "4. 现场支持：质保期内免费现场支持，不限次数"
    )

    doc.add_heading("6.2 运维团队", level=2)
    doc.add_paragraph(
        "项目验收后，我方将组建专项运维团队，配置如下：\n"
        "- 运维经理：1人\n"
        "- 运维工程师：2人\n"
        "- 数据库管理员：1人"
    )

    # 故意埋坑：等于没写
    doc.add_heading("6.3 培训方案", level=2)
    doc.add_paragraph("培训方案待定。")

    return doc


def main():
    output_dir = Path("data/input")
    output_dir.mkdir(parents=True, exist_ok=True)
    doc = create_test_bid_document()
    output_path = output_dir / "测试投标文件.docx"
    doc.save(str(output_path))
    print(f"✅ 测试投标文件已生成: {output_path}")
    print(f"   文件大小: {output_path.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    main()
