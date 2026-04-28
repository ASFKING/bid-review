# test_read_docx.py
# 这是我们的第一个 Python 脚本——读取 Word 文档并提取章节结构

from docx import Document  # 导入 python-docx 库，相当于 Java 的 import


def read_docx(file_path: str) -> dict:
    """
    读取 Word 文档，提取段落文本
    参数 file_path: Word 文件路径，类似 Java 的方法参数
    返回值: 一个字典，包含文档信息（Python 的 dict ≈ Java 的 Map）
    """
    # 创建 Document 对象，类似 Java 的 new Document(filePath)
    doc = Document(file_path)

    # 提取所有段落
    # Python 的列表推导式，等价于 Java 的 stream().map().collect()
    paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]

    # 构建结果字典（Python 的 dict ≈ Java 的 HashMap）
    result = {
        "total_paragraphs": len(paragraphs),      # 总段落数
        "paragraphs": paragraphs[:10],             # 先只看前 10 段（调试用）
        "total_chars": sum(len(p) for p in paragraphs),  # 总字符数
    }

    return result


# ===== 主程序入口 =====
# Python 的 if __name__ == "__main__" ≈ Java 的 public static void main
if __name__ == "__main__":
    # 先用一个测试文件试试
    # 把你的标书文件放到 data/input/ 目录下，然后修改文件名
    file_path = "data/input/test.docx"

    try:
        # try/except ≈ Java 的 try/catch
        info = read_docx(file_path)

        # print() ≈ Java 的 System.out.println()
        print(f"文档总段落数: {info['total_paragraphs']}")
        print(f"文档总字符数: {info['total_chars']}")
        print("\n--- 前 10 个段落 ---")
        for i, para in enumerate(info['paragraphs'], 1):
            print(f"  [{i}] {para[:80]}...")  # 每段只显示前 80 个字符

    except FileNotFoundError:
        print(f"文件不存在: {file_path}")
        print("请把标书文件放到 data/input/ 目录下")
    except Exception as e:
        print(f"读取出错: {e}")
